from broth import Broth
from csv import reader as csv_reader
from date_extractor import extract_date
from docx import Document
from io import BytesIO
from openpyxl import load_workbook
from os.path import isfile
from re import match
from requests import get
from validators import url

def clean(value):
    if isinstance(value,float) or isinstance(value,int):
        return value
    elif isinstance(value, str) or isinstance(value, unicode):
        value = value.strip().strip('"').strip('"').strip()
        date = extract_date(value)
        if date:
            return date
        elif match("^\d+?\.\d+?$", value):
            return float(value)
        elif value.isdigit():
            return int(value)
        else:
            return value

def extract_tables(data, debug=False):

    try:

        if debug: print "starting extract_tables"
        if debug: print "type of data:", type(data)

        data_type = str(type(data))


        unsupported_extensions = ("doc", "xls")

        # we are not using isinstance here because the user might not have Django installed
        if data_type in ("<type 'file'>", "<class 'django.core.files.uploadedfile.InMemoryUploadedFile'>"):
            filename = data.name
            extension = filename.split(".")[-1]
            if filename.endswith(".csv"):
                return extract_tables_from_csv(data)
            elif filename.endswith(".docx"):
                return extract_tables_from_doc(data)
            elif extension in ("xlsx", "xlsm", "xltx", "xltm"):
                return extract_tables_from_excel_spreadsheet(data)
            elif filename.endswith("pdf"):
                return extract_tables_from_pdf(data)
            elif filename.endswith("tsv"):
                return extract_tables_from_tsv(data)
            elif extension in unsupported_extensions:
                raise Exception("table-extractor does not the " + extension + " extension")

        elif isinstance(data, str) or isinstance(data, unicode):
            if isfile(data):
                with open(data) as f:
                    return extract_tables(f)
            elif url(data):
                extension = data.split(".")[-1]
                #if extension in ("csv", "tsv", "xls", "xlsm", "xlsx"):
                if extension == "csv":
                    return extract_tables_from_csv(get(data).iter_lines())
                elif extension == "docx":
                    return extract_tables_from_doc(BytesIO(get(data).content))
                elif extension == "tsv":
                    return extract_tables_from_tsv(get(data).iter_lines())
                elif extension in ("xlsx", "xlsm", "xltx", "xltm"):
                    return extract_tables_from_excel_spreadsheet(BytesIO(get(data).content))
                elif extension in unsupported_extensions:
                    raise Exception("table-extractor does not the " + extension + " extension")
                else:
                    print "if not a known ending, assume it's html"
                    print "should watch out for large files like videos"
                    return extract_tables_from_html(get(data).text)

            elif data.count("<") > 5 and data.count(">") > 5:
                return extract_tables_from_html(data)

        else:
            print "[table-extractor] unhandled file type"

    except Exception as e:
        print "[table-extractor] caught exception in extract_tables:", e

def extract_tables_from_excel_spreadsheet(excel_file):
    # can't do this all in one line with list comprehension because sometimes openpyxl and .xlsx adds a row at ends that's all Nones
    tables = []
    for sheet in load_workbook(excel_file):
        rows = []
        for row in sheet:
            values = [clean(cell.value) for cell in row] 
            set_of_values = set(values)
            if not(len(set_of_values) == 1 and set_of_values.pop() == None):
                rows.append(values)
        tables.append(rows)
    return tables
     

def extract_tables_from_csv(csvfile, delimiter=",", quotechar='"', debug=False):
    # assumes commas separated for now
    rows = []
    for row in csv_reader(csvfile, delimiter=delimiter, quotechar=quotechar):
        if debug: print "row:", row
        values = [clean(cell) for cell in row]
        if debug: print "values:", values
        set_of_values = set(values)
        if not(len(set_of_values) == 1 and set_of_values.pop() == None):
            rows.append(values)
    tables = [rows]
    if debug: print "[table-extractor] finishing extract_tables_from_csv with", tables
    return tables
        

def extract_tables_from_doc(doc_file):
    tables = []
    for table in Document(doc_file).tables:
        rows = []
        for row in rows:
            values = [clean(cell.text) for cell in row.cells]
            set_of_values = set(values)
            if not(len(set_of_values) == 1 and set_of_values.pop() == None):
                rows.append(values)
        tables.append(rows)
    return tables

def extract_tables_from_html(html):
    tables = []
    for table in Broth(html).tables:
        rows = []
        header = [clean(th.text) for th in table.select("thead tr th")]
        if header:
            rows.append(header)

        for row in table.select("tbody tr"):
            tds = [clean(td.text) for td in row.select("td")]
            if tds:
                rows.append(tds)
        tables.append(rows)
    return tables
        

def extract_tables_from_tsv(tsvfile):
    return extract_tables_from_csv(tsvfile, delimiter="\t")



